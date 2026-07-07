import os
import glob
import re
# pyrefly: ignore [missing-import]
from pypdf import PdfReader
from docx import Document

# ── Constants ──────────────────────────────────────────────────────────────────
CHUNK_SIZE    = 1500  # bigger chunks = more complete resume sections per chunk
CHUNK_OVERLAP = 200   # overlap to avoid cutting mid-sentence


# ── Text Extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    """Extract all text from a PDF file, page by page."""
    try:
        reader = PdfReader(file_path)
        text_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t and t.strip():
                text_parts.append(t.strip())
        return "\n\n".join(text_parts)
    except Exception as e:
        print(f"[chunk] Error reading PDF {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path: str) -> str:
    """Extract all text from a DOCX file including paragraphs and tables."""
    try:
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_txt = cell.text.strip()
                    if cell_txt:
                        row_text.append(cell_txt)
                if row_text:
                    text_parts.append(" | ".join(row_text))
        return "\n".join(text_parts)
    except Exception as e:
        print(f"[chunk] Error reading DOCX {file_path}: {e}")
        return ""

def extract_text(file_path: str) -> str:
    """Detect file extension and extract text accordingly."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception:
            return ""


# ── Name Extraction ────────────────────────────────────────────────────────────

def clean_filename_to_name(filename: str) -> str:
    """Extract a clean candidate name from the filename."""
    name_part = os.path.splitext(filename)[0]
    # If there is a ' - ', take the part after it which usually has the candidate name
    if " - " in name_part:
        name_part = name_part.split(" - ")[-1].strip()
    
    # Replace separators with spaces
    name_part = re.sub(r'[_#\-\.\(\)\/]', ' ', name_part)
    
    # Remove roll numbers / alphanumeric words (like 21H51A6740, ranisri21)
    words = name_part.split()
    cleaned_words = []
    
    ignore_words = {
        "resume", "updated", "groundtruth", "profile", "sample", "project",
        "present", "rag", "pdf", "docx", "new", "final", "latest", "cv",
        "b", "tech", "csd", "cse", "ece", "it", "mca", "btech", "mtech",
        "ug", "pg", "copied", "copy"
    }
    
    for w in words:
        w_lower = w.lower()
        
        # If it's a number (year, version, etc.) or in ignore words, skip
        if w.isdigit() or w_lower in ignore_words:
            continue
            
        # If it contains university roll number pattern (letters and digits, e.g. 21H51A6740)
        if re.search(r'\d', w) and re.search(r'[A-Za-z]', w):
            # If the letters part is long, we can keep the letters part
            letters = re.sub(r'\d+', '', w)
            if len(letters) >= 4 and letters.lower() not in ignore_words:
                cleaned_words.append(letters)
            continue
            
        cleaned_words.append(w)
        
    result = " ".join(cleaned_words).strip()
    result = result.title()
    return result if result else name_part.strip().title()



# ── Smart Chunking ─────────────────────────────────────────────────────────────

def split_text_into_chunks(
    text: str,
    chunk_size: int = CHUNK_SIZE,
    overlap: int = CHUNK_OVERLAP
) -> list[str]:
    """
    Splits text into overlapping chunks.
    Priority: paragraph boundary -> sentence boundary -> hard cut.
    Each step advances by at least (chunk_size - overlap) characters.
    """
    text = re.sub(r'\n{3,}', '\n\n', text).strip()

    # Short enough to keep as one chunk
    if len(text) <= chunk_size:
        return [text] if text else []

    chunks = []
    start  = 0
    step   = chunk_size - overlap   # minimum advance per iteration

    while start < len(text):
        end = start + chunk_size

        if end >= len(text):
            tail = text[start:].strip()
            if tail:
                chunks.append(tail)
            break

        # Try paragraph boundary (search in the latter half of the window)
        split_pos = text.rfind("\n\n", start + step, end)

        # Fall back to sentence boundary
        if split_pos < start + step:
            for sep in (". ", "? ", "! ", "\n"):
                pos = text.rfind(sep, start + step, end)
                if pos >= start + step:
                    split_pos = pos + len(sep)
                    break
            else:
                # Hard cut at end of window
                split_pos = end

        piece = text[start:split_pos].strip()
        if piece:
            chunks.append(piece)

        # Advance: next chunk starts (split_pos - overlap) but at least step chars forward
        start = max(start + step, split_pos - overlap)

    return chunks


# ── Main Entry ─────────────────────────────────────────────────────────────────

def process_resumes_folder(folder_path: str = "resumes") -> list[dict]:
    """
    Scans the folder, reads all PDF and DOCX files.
    Each resume → multiple overlapping chunks (1500 chars, 200 overlap).

    IMPORTANT: chunk text contains ONLY the resume content (no metadata prefix),
    so the embedding purely represents resume content for accurate similarity search.
    Candidate name is stored in metadata for context building.

    Returns list of dicts: [{'text', 'source', 'chunk_id', 'candidate'}]
    """
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"Resumes folder not found: {os.path.abspath(folder_path)}")

    all_chunks: list[dict] = []
    patterns = [
        os.path.join(folder_path, "*.pdf"),
        os.path.join(folder_path, "*.docx"),
        os.path.join(folder_path, "*.PDF"),
        os.path.join(folder_path, "*.DOCX"),
    ]

    files: list[str] = []
    for pat in patterns:
        files.extend(glob.glob(pat))

    # De-duplicate on Windows (case-insensitive paths)
    files = list({os.path.abspath(f).lower(): os.path.abspath(f) for f in files}.values())
    files.sort()

    total_chunks = 0
    for file_path in files:
        filename  = os.path.basename(file_path)
        candidate = clean_filename_to_name(filename)

        text = extract_text(file_path)

        # ── Scanned / unreadable fallback ──
        if not text.strip():
            fallback_text = (
                f"Candidate Name: {candidate}\n"
                f"File: {filename}\n"
                f"Note: This resume could not be parsed (possibly a scanned image PDF)."
            )
            all_chunks.append({
                "text":      fallback_text,
                "source":    filename,
                "chunk_id":  f"{filename}__chunk_0",
                "candidate": candidate,
            })
            total_chunks += 1
            print(f"[chunk] {filename}: 1 fallback chunk (unreadable)")
            continue

        # ── Split into content chunks ──
        sub_chunks = split_text_into_chunks(text)
        if not sub_chunks:
            sub_chunks = [text.strip()]

        for i, sub in enumerate(sub_chunks):
            # Prepend candidate name and source info to the chunk text so the embedding captures it
            # and the retriever matches when searching for that candidate.
            chunk_text = f"Candidate: {candidate}\nSource: {filename}\n\n{sub}"
            all_chunks.append({
                "text":      chunk_text,
                "source":    filename,
                "chunk_id":  f"{filename}__chunk_{i}",
                "candidate": candidate,
            })

        total_chunks += len(sub_chunks)
        print(f"[chunk] {filename}: {len(sub_chunks)} chunks  (candidate: {candidate})")

    print(f"\n[chunk] [OK] {len(files)} resumes -> {total_chunks} total chunks\n")
    return all_chunks
